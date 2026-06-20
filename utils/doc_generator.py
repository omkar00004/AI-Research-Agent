"""Professional DOCX report generator with styled formatting.

Renders mermaid diagrams as images via the mermaid.ink API.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import base64
import io
import re
import tempfile
import urllib.request
import urllib.error


ACCENT = RGBColor(0x1B, 0x4F, 0x8A)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "D6E4F7"


def set_cell_bg(cell, hex_color):
    """Apply a background colour to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_border_bottom(paragraph, color="1B4F8A", size=12):
    """Add a coloured bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def render_mermaid_to_png(mermaid_code: str) -> bytes | None:
    """Render mermaid code to PNG bytes using the mermaid.ink API.

    Returns PNG bytes on success, or None on failure.
    """
    try:
        # Base64-encode the mermaid code for the mermaid.ink URL
        encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii")
        url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white"

        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Atlas-Research-Agent/1.0"},
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            return resp.read()
    except Exception as e:
        print(f"[doc_generator] mermaid.ink render failed: {e}")
        return None


def add_code_block(doc, code_text: str, label: str = "Diagram Code"):
    """Add a styled code block as a fallback for mermaid diagrams."""
    # Label paragraph
    label_para = doc.add_paragraph()
    label_run = label_para.add_run(f"[{label}]")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = GRAY
    label_run.font.name = "Arial"

    # Code content in a shaded paragraph
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code_text)
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(8)
    code_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Add background shading to the paragraph
    pPr = code_para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)


def generate_docx(topic: str, report_text: str, sources: list) -> bytes:
    """Generate a professional .docx report from the markdown report text.

    Detects mermaid code blocks and renders them as images.
    Returns the document as bytes ready for download.
    """

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # ---- HEADER BLOCK ----
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(topic.upper())
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = ACCENT
    title_run.font.name = "Arial"
    title_para.paragraph_format.space_after = Pt(4)
    add_border_bottom(title_para)

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(
        f"AI Multi-Agent Research Report  |  Generated: {datetime.now().strftime('%B %d, %Y')}  |  Sources: {len(sources)}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = GRAY
    meta_run.font.name = "Arial"
    meta_para.paragraph_format.space_after = Pt(16)

    # ---- PARSE AND RENDER REPORT ----
    lines = report_text.strip().split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # ---- Mermaid code block detection ----
        if stripped.startswith("```mermaid"):
            # Collect all lines until closing ```
            mermaid_lines = []
            i += 1
            while i < len(lines):
                if lines[i].strip() == "```":
                    i += 1
                    break
                mermaid_lines.append(lines[i])
                i += 1

            mermaid_code = "\n".join(mermaid_lines).strip()
            if mermaid_code:
                # Try rendering via mermaid.ink
                png_bytes = render_mermaid_to_png(mermaid_code)
                if png_bytes:
                    # Embed image in doc
                    img_stream = io.BytesIO(png_bytes)
                    try:
                        doc.add_picture(img_stream, width=Inches(6))
                        # Center the image
                        last_para = doc.paragraphs[-1]
                        last_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        last_para.paragraph_format.space_before = Pt(8)
                        last_para.paragraph_format.space_after = Pt(8)
                    except Exception:
                        add_code_block(doc, mermaid_code, "Mermaid Diagram")
                else:
                    add_code_block(doc, mermaid_code, "Mermaid Diagram")
            continue

        # ---- Other code block (skip raw code fences) ----
        if stripped.startswith("```"):
            # Non-mermaid code block — collect and add as code
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines):
                if lines[i].strip() == "```":
                    i += 1
                    break
                code_lines.append(lines[i])
                i += 1

            code_text = "\n".join(code_lines).strip()
            if code_text:
                add_code_block(doc, code_text, lang or "Code")
            continue

        # Section header (###)
        if stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            h = doc.add_paragraph()
            run = h.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = ACCENT
            run.font.name = "Arial"
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            add_border_bottom(h, color="D6E4F7", size=6)

        # H2 (##)
        elif stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            h = doc.add_paragraph()
            run = h.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = ACCENT
            run.font.name = "Arial"
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6)
            add_border_bottom(h)

        # H1 (#)
        elif stripped.startswith("# "):
            heading_text = stripped[2:].strip()
            h = doc.add_paragraph()
            run = h.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = ACCENT
            run.font.name = "Arial"
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            add_border_bottom(h)

        # Blockquotes
        elif stripped.startswith("> "):
            quote_text = stripped[2:].strip()
            # Remove surrounding quotes if present
            quote_text = quote_text.strip('"').strip('"').strip('"')
            p = doc.add_paragraph()
            run = p.add_run(f'"{quote_text}"')
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            run.font.name = "Arial"
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add left border styling via shading
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "16")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "3B82F6")
            pBdr.append(left)
            pPr.append(pBdr)

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            # Clean bold markers
            parts = re.split(r'\*\*(.*?)\*\*', bullet_text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                run.bold = (j % 2 == 1)
                run.font.name = "Arial"
                run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        # Numbered list
        elif re.match(r'^\d+\.', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                run.bold = (j % 2 == 1)
                run.font.name = "Arial"
                run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            add_border_bottom(p, color="CCCCCC", size=4)

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            # Handle inline bold
            parts = re.split(r'\*\*(.*?)\*\*', stripped)
            for j, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.bold = (j % 2 == 1)
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        i += 1

    # ---- SOURCES TABLE ----
    if sources:
        doc.add_paragraph()
        src_heading = doc.add_paragraph()
        run = src_heading.add_run("SOURCES & REFERENCES")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = ACCENT
        run.font.name = "Arial"
        src_heading.paragraph_format.space_before = Pt(16)
        src_heading.paragraph_format.space_after = Pt(8)
        add_border_bottom(src_heading)

        # Deduplicate sources
        seen_urls = set()
        unique_sources = []
        for s in sources:
            if s.get("url") not in seen_urls:
                seen_urls.add(s.get("url"))
                unique_sources.append(s)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        for cell, text in zip(header_cells, ["Title", "URL"]):
            set_cell_bg(cell, "1B4F8A")
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "Arial"
            run.font.size = Pt(10)

        # Source rows
        for idx, source in enumerate(unique_sources[:20]):
            row = table.add_row()
            if idx % 2 == 0:
                set_cell_bg(row.cells[0], "F5F8FC")
                set_cell_bg(row.cells[1], "F5F8FC")

            title_run = row.cells[0].paragraphs[0].add_run(source.get("title", "")[:80])
            title_run.font.name = "Arial"
            title_run.font.size = Pt(9)

            url_run = row.cells[1].paragraphs[0].add_run(source.get("url", "")[:80])
            url_run.font.name = "Arial"
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = ACCENT

        # Column widths
        for row in table.rows:
            row.cells[0].width = Inches(3.0)
            row.cells[1].width = Inches(4.5)

    # ---- FOOTER ----
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"Generated by AI Multi-Agent Research System  |  {datetime.now().strftime('%Y')}  |  Powered by LangGraph + Groq"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = GRAY
    footer_run.font.name = "Arial"
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_bottom(footer_para, color="CCCCCC", size=4)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
